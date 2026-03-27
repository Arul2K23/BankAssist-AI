import os
import uuid
import json
from typing import List, Dict, Any, Optional
from langchain_core.documents import Document
from langchain_community.document_loaders import (
    PyPDFLoader,
    TextLoader,
    CSVLoader,
    UnstructuredExcelLoader,
    UnstructuredPowerPointLoader,
)
from langchain_text_splitters import RecursiveCharacterTextSplitter
from app.core.config import settings
from app.models.auth import AccessLevel

# ── Stable model lists ────────────────────────────────────────────────────────
GEMINI_MODELS = ["gemini-2.0-flash", "gemini-1.5-flash"]
GROQ_MODELS   = ["llama-3.3-70b-versatile", "llama3-8b-8192"]


# ── Custom DOCX loader ────────────────────────────────────────────────────────
def load_docx(file_path: str) -> List[Document]:
    """
    Walks the document body in document order — paragraphs and tables together —
    so that table cells are never silently skipped.

    ROOT CAUSE OF THE BUG THIS FIXES:
    LangChain's UnstructuredWordDocumentLoader iterates doc.paragraphs which
    only contains text that lives OUTSIDE of tables. In a .docx file, table
    cell text lives in a completely separate XML tree (<w:tbl>). This means
    every table in every document was completely invisible to the RAG system —
    including branch working hours, limits, charges, etc.

    This loader walks doc.element.body directly, which contains both <w:p>
    (paragraph) and <w:tbl> (table) elements in their original document order.
    Each table row is rendered as a markdown-style pipe-delimited line so the
    LLM can understand the structure.
    """
    from docx import Document as DocxDocument

    doc = DocxDocument(file_path)
    blocks: List[str] = []

    def parse_table(table) -> str:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")
        # Insert a separator after the header row
        if len(rows) > 1:
            rows.insert(1, "|" + "|".join(["---"] * len(table.columns)) + "|")
        return "\n".join(rows)

    # Walk raw XML body to preserve document order
    for child in doc.element.body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "p":
            text = "".join(
                node.text or ""
                for node in child.iter()
                if node.tag.endswith("}t")
            ).strip()
            if text:
                blocks.append(text)

        elif tag == "tbl":
            for table in doc.tables:
                if table._tbl is child:
                    table_text = parse_table(table)
                    if table_text.strip():
                        blocks.append(table_text)
                    break

    full_text = "\n\n".join(blocks)
    return [Document(page_content=full_text, metadata={"source": file_path})]


class IngestionService:
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.CHUNK_SIZE,
            chunk_overlap=settings.CHUNK_OVERLAP,
            separators=["\n\n", "\n", ".", " ", ""],
        )
        self._gemini_idx = 0
        self._groq_idx   = 0
        self._llm        = None
        self._provider   = None
        self._init_llm()

    # ── LLM setup ─────────────────────────────────────────────────────────────

    def _init_llm(self):
        groq_key = getattr(settings, "GROQ_API_KEY", None)
        if groq_key:
            try:
                from langchain_groq import ChatGroq
                self._llm = ChatGroq(
                    model=GROQ_MODELS[self._groq_idx],
                    api_key=groq_key,
                    temperature=0,
                    max_retries=1,
                )
                self._provider = "groq"
                print(f"DEBUG INGESTION: Using Groq ({GROQ_MODELS[self._groq_idx]})")
                return
            except ImportError:
                print("WARNING INGESTION: langchain-groq not installed — falling back to Gemini.")
        self._init_gemini()

    def _init_gemini(self, idx: Optional[int] = None):
        from langchain_google_genai import ChatGoogleGenerativeAI
        if idx is not None:
            self._gemini_idx = idx
        name = GEMINI_MODELS[self._gemini_idx]
        self._llm = ChatGoogleGenerativeAI(
            model=name,
            google_api_key=settings.GEMINI_API_KEY,
            temperature=0,
            max_retries=1,
        )
        self._provider = "gemini"
        print(f"DEBUG INGESTION: Using Gemini ({name})")

    def _rotate_llm(self):
        if self._provider == "groq":
            next_idx = self._groq_idx + 1
            if next_idx < len(GROQ_MODELS):
                self._groq_idx = next_idx
                from langchain_groq import ChatGroq
                self._llm = ChatGroq(
                    model=GROQ_MODELS[self._groq_idx],
                    api_key=getattr(settings, "GROQ_API_KEY", None),
                    temperature=0,
                    max_retries=1,
                )
                print(f"DEBUG INGESTION: Rotated Groq → {GROQ_MODELS[self._groq_idx]}")
            else:
                print("DEBUG INGESTION: Groq exhausted — switching to Gemini.")
                self._init_gemini(0)
        else:
            next_idx = self._gemini_idx + 1
            if next_idx < len(GEMINI_MODELS):
                self._init_gemini(next_idx)
            else:
                print("WARNING INGESTION: All providers exhausted.")

    @staticmethod
    def _extract_text(raw_content) -> str:
        if isinstance(raw_content, list):
            return raw_content[0].get("text", str(raw_content)) if raw_content else ""
        if isinstance(raw_content, str):
            try:
                parsed = json.loads(raw_content)
                if isinstance(parsed, list):
                    return parsed[0].get("text", raw_content) if parsed else raw_content
            except json.JSONDecodeError:
                pass
            return raw_content
        return str(raw_content)

    @staticmethod
    def _is_retryable(err: str) -> bool:
        return any(code in err for code in ("429", "404", "503")) or "quota" in err.lower()

    def _invoke_with_rotation(self, prompt_or_messages, max_attempts: int = 4) -> str:
        for attempt in range(max_attempts):
            try:
                response = self._llm.invoke(prompt_or_messages)
                return self._extract_text(response.content)
            except Exception as e:
                err = str(e)
                if self._is_retryable(err):
                    print(f"DEBUG INGESTION: attempt {attempt + 1} failed ({err[:60]}) — rotating…")
                    self._rotate_llm()
                    continue
                raise
        raise RuntimeError("All ingestion LLM attempts exhausted.")

    # ── Summarisation ─────────────────────────────────────────────────────────

    def generate_summary(self, text: str) -> str:
        sample = text[:5000]
        prompt = (
            "Summarize the following document text in exactly 2-3 concise bullet points "
            "for an enterprise database overview:\n\n" + sample
        )
        try:
            return self._invoke_with_rotation(prompt)
        except Exception as e:
            print(f"Summarisation failed: {e}")
            return "Summary generation unavailable."

    # ── Image extraction ──────────────────────────────────────────────────────

    def _process_image(self, file_path: str) -> List[Document]:
        import base64
        from langchain_core.messages import HumanMessage

        with open(file_path, "rb") as f:
            image_b64 = base64.b64encode(f.read()).decode("utf-8")

        ext = os.path.splitext(file_path)[1].lower().lstrip(".")
        if ext == "jpg":
            ext = "jpeg"

        message = HumanMessage(
            content=[
                {
                    "type": "text",
                    "text": (
                        "Extract all text from this image perfectly. "
                        "If there are charts, tables, or diagrams, describe them in "
                        "immense detail so they can be indexed in a database for searching. "
                        "Start directly with the extracted text/description."
                    ),
                },
                {
                    "type": "image_url",
                    "image_url": {"url": f"data:image/{ext};base64,{image_b64}"},
                },
            ]
        )

        try:
            text_content = self._invoke_with_rotation([message])
            return [Document(page_content=text_content, metadata={"source": file_path})]
        except Exception as e:
            print(f"Failed to process image {file_path}: {e}")
            return [
                Document(
                    page_content="Error processing image — all providers exhausted.",
                    metadata={"source": file_path},
                )
            ]

    # ── Document loading ──────────────────────────────────────────────────────

    def load_document(self, file_path: str) -> List[Document]:
        ext = os.path.splitext(file_path)[1].lower()
        try:
            if ext == ".pdf":
                return PyPDFLoader(file_path).load()

            elif ext in (".doc", ".docx"):
                # ✅ Custom loader — preserves table content
                # UnstructuredWordDocumentLoader silently drops all tables
                return load_docx(file_path)

            elif ext == ".csv":
                return CSVLoader(file_path).load()

            elif ext in (".xls", ".xlsx"):
                return UnstructuredExcelLoader(file_path).load()

            elif ext in (".ppt", ".pptx"):
                return UnstructuredPowerPointLoader(file_path).load()

            elif ext in (".png", ".jpg", ".jpeg"):
                return self._process_image(file_path)

            else:
                return TextLoader(file_path).load()

        except Exception as e:
            print(f"Failed to load {file_path}: {e} — returning error doc.")
            return [
                Document(
                    page_content=f"Error reading file: {e}",
                    metadata={"source": file_path},
                )
            ]

    # ── Chunking + metadata enrichment ────────────────────────────────────────

    def process_document(
        self, file_path: str, metadata: Dict[str, Any]
    ) -> List[Dict[str, Any]]:
        docs   = self.load_document(file_path)
        chunks = self.text_splitter.split_documents(docs)

        processed = []
        for i, chunk in enumerate(chunks):
            chunk_meta = {
                "doc_id":        str(uuid.uuid4()),
                "parent_doc_id": metadata.get("parent_doc_id"),
                "source":        os.path.basename(file_path),
                "chunk_index":   i,
                "access_level":  metadata.get("access_level", AccessLevel.INTERNAL),
                "department":    metadata.get("department", "General"),
                "version":       metadata.get("version", "1.0"),
                "timestamp":     metadata.get("timestamp", ""),
            }
            processed.append({"text": chunk.page_content, "metadata": chunk_meta})

        print(
            f"DEBUG INGESTION: {len(processed)} chunks from "
            f"{os.path.basename(file_path)}"
        )
        return processed


ingestion_service = IngestionService()